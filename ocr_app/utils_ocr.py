import pytesseract
import cv2
import numpy as np
from PIL import Image
from docx import Document as DocxDocument
import fitz  # PyMuPDF
import re
from urllib.parse import unquote
from io import BytesIO
from typing import Union, BinaryIO

# ---------------------------------------------------------------------------
# Helper: accept either path str or file-like object and return PIL Image
# ---------------------------------------------------------------------------

def _load_image(img_source: Union[str, BinaryIO]):
    """Return cv2 image from a file path or a binary stream."""
    try:
        if isinstance(img_source, str):
            img = cv2.imread(img_source)
        else:
            # Stream → bytes → numpy array → cv2 image
            img_bytes = img_source.read()
            np_arr = np.frombuffer(img_bytes, np.uint8)
            img = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
        return img
    except Exception:
        return None

# ---------------------------------------------------------------------------
# OCR FOR IMAGES (Arabic + English)
# ---------------------------------------------------------------------------

def extract_text_from_image(image_source: Union[str, BinaryIO]):
    """Enhanced OCR for Arabic/English images. Accepts path or stream."""
    try:
        img = _load_image(image_source)
        if img is None:
            return ""

        # --- Pre‑processing (lab, clahe, thresh) ---
        lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        cl = clahe.apply(l)
        limg = cv2.merge((cl, a, b))
        enhanced = cv2.cvtColor(limg, cv2.COLOR_LAB2BGR)
        gray = cv2.cvtColor(enhanced, cv2.COLOR_BGR2GRAY)
        thresh = cv2.adaptiveThreshold(gray, 255,
                                       cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                       cv2.THRESH_BINARY, 41, 2)
        cleaned = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, np.ones((1, 1), np.uint8))

        custom = r"-l ara+eng --psm 6 -c preserve_interword_spaces=1"
        return pytesseract.image_to_string(cleaned, config=custom)
    except Exception as e:
        print(f"OCR Image Error: {e}")
        return ""

# ---------------------------------------------------------------------------
# OCR FOR PDF (Arabic preservation) – accepts path or stream
# ---------------------------------------------------------------------------

def _open_pdf(pdf_source: Union[str, BinaryIO]):
    if isinstance(pdf_source, str):
        return fitz.open(pdf_source)
    # For stream use bytes → PyMuPDF via `fitz.open(stream=..., filetype='pdf')`
    pdf_source.seek(0)
    return fitz.open(stream=pdf_source.read(), filetype="pdf")


def extract_text_from_pdf(pdf_source: Union[str, BinaryIO]):
    """Extract text from PDF preserving ligatures (Arabic)."""
    try:
        doc = _open_pdf(pdf_source)
        flags = fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE
        text = "\n".join(page.get_text("text", flags=flags) for page in doc)
        return text
    except Exception as e:
        print(f"PDF Extraction Error: {e}")
        return ""

# ---------------------------------------------------------------------------
# OCR FOR WORD (docx) – accepts path or stream
# ---------------------------------------------------------------------------

def extract_text_from_word(word_source: Union[str, BinaryIO]):
    try:
        if isinstance(word_source, str):
            doc = DocxDocument(word_source)
        else:
            word_source.seek(0)
            doc = DocxDocument(BytesIO(word_source.read()))
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        print(f"Word Extraction Error: {e}")
        return ""

# ---------------------------------------------------------------------------
# ADVANCED NORMALIZATION FUNCTION – unchanged but wrapped in function
# ---------------------------------------------------------------------------

def advanced_normalize_text(text: str):
    """Arabic text normalization with special character preservation."""
    if not text:
        return ""

    for _ in range(3):
        try:
            if '%' in text:
                decoded = unquote(text)
                if decoded != text:
                    text = decoded
                else:
                    break
        except Exception:
            break

    text = re.sub(r'[\u064B-\u0652\u0670\u0640\u06D6-\u06ED\u08F0-\u08FF]', '', text)

    replacements = {
        'أ': 'ا', 'إ': 'ا', 'آ': 'ا', 'ٱ': 'ا',
        'ة': 'ه', 'ت': 'ت',
        'ى': 'ي', 'ئ': 'ي', 'ؤ': 'و',
        'ه': 'ه', 'ة': 'ه',
        'ك': 'ك', 'ی': 'ي', 'ے': 'ي',
        '\u200c': '', '\u200d': '', '\u200e': '', '\u200f': '',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    text = text.lower()
    text = re.sub(r'[\|/\\]+', ' ', text)
    text = re.sub(r'[^\w\s\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF_-]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text
